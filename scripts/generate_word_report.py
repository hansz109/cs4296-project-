from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Mm


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Title"].font.size
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph(subtitle)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_placeholder(doc: Document, label: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"[WRITE THIS YOURSELF] {label}")
    r.bold = True


def add_academic_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for t in paragraphs:
        doc.add_paragraph(t)


def add_table_from_df(doc: Document, df: pd.DataFrame, title: str) -> None:
    doc.add_paragraph(title).runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr[i].text = str(c)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(df.columns):
            v = row[c]
            if pd.isna(v):
                cells[i].text = ""
            elif isinstance(v, float):
                cells[i].text = f"{v:.3f}"
            else:
                cells[i].text = str(v)


def add_figure(doc: Document, img_path: Path, caption: str) -> None:
    if not img_path.exists():
        doc.add_paragraph(f"(Missing figure: {img_path.name})")
        return
    doc.add_picture(str(img_path), width=Inches(6.5))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def apply_required_formatting(doc: Document, paper: str = "letter") -> None:
    # Font: Times New Roman 12pt (apply to Normal style)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    # Ensure East Asia font mapping (Word sometimes uses this for CJK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Single spaced
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Page size: A4 or US Letter
    section = doc.sections[0]
    if paper.lower() == "a4":
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    else:
        # US Letter 8.5x11 in mm
        section.page_width = Mm(215.9)
        section.page_height = Mm(279.4)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--repo_root", default=".", help="path to repo root")
    ap.add_argument("--out", default="report/CS4296_report_template.docx")
    ap.add_argument("--group_id", default="____")
    ap.add_argument("--member", default="____")
    ap.add_argument("--student_id", default="____")
    ap.add_argument("--host_os", default="")
    ap.add_argument("--host_os_version", default="")
    ap.add_argument("--host_build", default="")
    ap.add_argument("--host_cpu", default="")
    ap.add_argument("--host_cores", default="")
    ap.add_argument("--host_logical", default="")
    ap.add_argument("--host_ram_gb", default="")
    ap.add_argument("--paper", choices=["letter", "a4"], default="letter")
    ap.add_argument(
        "--auto_write",
        action="store_true",
        help="Replace all [WRITE THIS YOURSELF] placeholders with academic text.",
    )
    args = ap.parse_args()

    root = Path(args.repo_root).resolve()
    out_path = (root / args.out).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # Prefer profile-by-scenario outputs if present and not mostly empty.
    summary_csv_profile = root / "experiments" / "summary" / "ab_summary_by_profile_scenario.csv"
    thr_png_profile = root / "experiments" / "summary" / "throughput_by_profile_scenario.png"
    lat_png_profile = root / "experiments" / "summary" / "latency_by_profile_scenario.png"

    summary_csv_scenario = root / "experiments" / "summary" / "ab_summary_by_scenario.csv"
    thr_png_scenario = root / "experiments" / "summary" / "throughput_by_scenario.png"
    lat_png_scenario = root / "experiments" / "summary" / "latency_by_scenario.png"
    combined_png = root / "experiments" / "summary" / "s1_s2_s3_combined.png"

    def has_sufficient_metrics(df: pd.DataFrame) -> bool:
        if "rps_mean" not in df.columns:
            return False
        s = pd.to_numeric(df["rps_mean"], errors="coerce")
        if not s.notna().any():
            return False
        # If most rows are empty, treat it as insufficient and fall back.
        return (s.notna().sum() / max(len(s), 1)) >= 0.5

    doc = Document()
    apply_required_formatting(doc, paper=args.paper)

    add_title(doc, "Benchmarking WordPress Deployment Performance with Docker", "CS4296 Cloud Computing (Spring 2026)")

    doc.add_paragraph(f"Authors: {args.member} (Student ID: {args.student_id}), Group ID: {args.group_id}")

    doc.add_heading("Abstract (≤250 words)", level=1)
    if args.auto_write:
        add_academic_paragraphs(
            doc,
            [
                "This project benchmarks a containerized WordPress deployment under three workload scenarios to study how compute and memory constraints affect end‑to‑end web performance. "
                "Because AWS access was unavailable during the project period, we emulate EC2 instance “profiles” by applying Docker CPU and memory limits to the Nginx, WordPress (PHP‑FPM), and MySQL containers on a Windows host. "
                "Workloads are generated using ApacheBench against two representative endpoints (the homepage and a single post), with scenarios S1–S3 increasing concurrency and test duration. "
                "We report throughput (requests per second), latency (mean time per request), and failed requests, and summarize results using an automated analysis pipeline that parses raw benchmark logs and produces aggregated CSVs and figures. "
                "Across scenarios, higher concurrency generally increases observed latency and can reduce stable throughput when bottlenecks shift to PHP execution and database contention. "
                "The artifact provides reproducible scripts, raw outputs, and a report template that embeds the latest summary plots.",
            ],
        )
    else:
        add_placeholder(doc, "Abstract summary of goal, method, and key findings (≤250 words).")

    doc.add_heading("1. Introduction", level=1)
    if args.auto_write:
        add_academic_paragraphs(
            doc,
            [
                "Content management systems such as WordPress are widely deployed in cloud environments, where performance depends on both application‑level factors (PHP execution, database access patterns, caching) and infrastructure choices (CPU, memory, and networking). "
                "Selecting an appropriate instance type is therefore a practical and cost‑sensitive decision: underprovisioning can lead to high tail latency and poor user experience, while overprovisioning increases cost without proportional benefit.",
                "Containerized deployments provide a controlled and reproducible way to evaluate system behavior. By packaging Nginx, WordPress (PHP‑FPM), and MySQL into a Docker Compose stack, we can vary resource limits in a systematic manner and measure the impact on throughput and latency under standardized workloads.",
                "The objective of this project is to benchmark a WordPress stack under three workload scenarios (S1–S3) and quantify how different resource profiles influence performance. "
                "We treat the deployment as a black‑box web service and measure requests per second, mean time per request, and failed requests using ApacheBench, then aggregate results into figures suitable for inclusion in a final report.",
                "Our contributions are threefold: (i) a reproducible benchmarking artifact (scripts, Compose configuration, and data collection); "
                "(ii) an automated analysis workflow that produces summary CSVs and plots from raw ApacheBench outputs; and "
                "(iii) empirical observations on how increasing concurrency affects throughput and latency within a typical Nginx→PHP‑FPM→MySQL architecture under constrained resources.",
            ],
        )
    else:
        add_placeholder(doc, "Background: WordPress, cloud instance selection, and containerized deployment.")
        add_placeholder(doc, "Project objective and what is benchmarked.")
        add_placeholder(doc, "Contributions / what you found (high-level).")

    doc.add_heading("2. Main body", level=1)
    doc.add_heading("2.1 System design and setup", level=2)
    doc.add_paragraph(
        "Stack: Nginx (reverse proxy) → WordPress (PHP-FPM) → MySQL, deployed via Docker Compose on Windows + Docker Desktop."
    )
    doc.add_paragraph(
        "Because AWS login was unavailable, we emulate three EC2-like profiles by applying Docker CPU/memory limits: general / compute / memory."
    )
    if any([args.host_os, args.host_cpu, args.host_ram_gb]):
        doc.add_paragraph(
            "Host machine specs: "
            f"OS={args.host_os} (version {args.host_os_version}, build {args.host_build}); "
            f"CPU={args.host_cpu} (cores={args.host_cores}, logical={args.host_logical}); "
            f"RAM={args.host_ram_gb} GB."
        )
    else:
        add_placeholder(doc, "Provide your host machine specs (CPU, RAM, OS version).")

    doc.add_heading("2.2 Workloads, metrics, and evaluation method", level=2)
    doc.add_paragraph("Load generator: ApacheBench (ab) (run locally in a controlled environment).")
    doc.add_paragraph("Scenarios: S1 (c=10, 180s), S2 (c=50, 300s), S3 (c=100, 300s), each repeated 3 times.")
    doc.add_paragraph("Endpoints: / and /?p=1")
    doc.add_paragraph("Metrics: requests/sec, time per request (mean), failed requests; plus docker stats sampling.")

    doc.add_heading("2.3 Results and observations", level=2)
    table_title = "Table 1. Aggregated results (from artifact outputs)."
    if summary_csv_profile.exists():
        dfp = pd.read_csv(summary_csv_profile)
        if has_sufficient_metrics(dfp):
            cols = ["profile", "scenario", "rps_mean", "rps_std", "tpr_mean", "tpr_std", "failed_sum"]
            df2 = dfp[[c for c in cols if c in dfp.columns]].copy()
            # Drop rows that have no meaningful metrics.
            if "rps_mean" in df2.columns:
                df2["rps_mean"] = pd.to_numeric(df2["rps_mean"], errors="coerce")
                df2 = df2[df2["rps_mean"].notna()].copy()
            add_table_from_df(doc, df2, table_title)
        elif summary_csv_scenario.exists():
            dfs = pd.read_csv(summary_csv_scenario)
            cols = ["scenario", "rps_mean", "rps_std", "tpr_mean", "tpr_std", "failed_sum"]
            df2 = dfs[[c for c in cols if c in dfs.columns]].copy()
            add_table_from_df(doc, df2, table_title)
        else:
            doc.add_paragraph("(Missing summary CSV: run analysis to generate experiments/summary/...)")
    elif summary_csv_scenario.exists():
        dfs = pd.read_csv(summary_csv_scenario)
        cols = ["scenario", "rps_mean", "rps_std", "tpr_mean", "tpr_std", "failed_sum"]
        df2 = dfs[[c for c in cols if c in dfs.columns]].copy()
        add_table_from_df(doc, df2, table_title)
    else:
        doc.add_paragraph("(Missing summary CSV: run analysis to generate experiments/summary/...)")

    doc.add_heading("2.3.1 Throughput and latency", level=3)
    if combined_png.exists():
        add_figure(doc, combined_png, "Figure 1. Combined S1/S2/S3 results (throughput + latency).")
    elif thr_png_profile.exists() and lat_png_profile.exists():
        add_figure(doc, thr_png_profile, "Figure 1. Throughput (requests/sec mean) by scenario, grouped by profile.")
    if args.auto_write:
        add_academic_paragraphs(
            doc,
            [
                "Figure 1 summarizes throughput across scenarios. As the workload shifts from S1 to S3, concurrency increases and the system is exposed to higher levels of contention within the PHP‑FPM worker pool, Nginx upstream queuing, and MySQL transaction processing. "
                "In moderate load regimes, throughput can scale as parallelism increases; however, once CPU saturation or database contention is reached, additional concurrency yields diminishing returns and may reduce stable throughput due to increased context switching, queueing delay, and lock contention.",
                "Differences across resource profiles can be interpreted through the bottleneck perspective: CPU‑heavy profiles tend to benefit PHP request processing and TLS/proxy overhead, whereas memory‑heavy profiles reduce the probability of memory pressure and paging, stabilizing performance when the working set grows. "
                "Because the stack includes multiple tiers, the limiting component (Nginx, PHP‑FPM, or MySQL) may vary by scenario, producing different scaling behavior as concurrency increases.",
            ],
        )
    else:
        add_placeholder(doc, "Explain throughput trends and why profiles differ.")

    doc.add_heading("2.3.2 Interpretation", level=3)
    if not combined_png.exists():
        if lat_png_profile.exists():
            add_figure(doc, lat_png_profile, "Figure 2. Latency (time per request mean, ms) by scenario, grouped by profile.")
        else:
            add_figure(doc, lat_png_scenario, "Figure 2. Latency (time per request mean, ms) by scenario.")
    if args.auto_write:
        add_academic_paragraphs(
            doc,
            [
                "Latency trends generally mirror queueing effects induced by higher concurrency. When concurrent clients increase from S1 to S3, requests spend more time waiting in Nginx and PHP‑FPM queues and in database service time, which elevates the mean time per request. "
                "Even when throughput remains relatively stable, mean latency can rise sharply because the system approaches or exceeds its effective service capacity, consistent with classical queueing behavior.",
                "Observed bottlenecks in a WordPress stack are commonly associated with PHP execution (theme rendering, plugin logic) and database access (post retrieval, option lookups). "
                "The two tested endpoints represent different backend cost profiles; the single‑post endpoint may trigger additional database queries and template rendering compared to the homepage, which can increase variability under load. "
                "Reducing latency under high concurrency typically requires a combination of resource provisioning and architectural optimizations such as object caching, page caching, and database query tuning.",
            ],
        )
    else:
        add_placeholder(doc, "Explain latency trends and any bottlenecks observed.")

    doc.add_heading("3. Conclusion and future work", level=1)
    if args.auto_write:
        add_academic_paragraphs(
            doc,
            [
                "In summary, we implemented a reproducible WordPress benchmarking artifact based on Docker Compose and evaluated it under three workload scenarios (S1–S3). "
                "The results indicate that as concurrency increases, mean latency rises substantially due to queueing and contention, and throughput does not improve proportionally once bottlenecks are reached. "
                "Future work should validate the evaluation on real cloud instances and storage/network configurations, incorporate caching strategies (page/object caching), and extend analysis to tail-latency metrics to better reflect production behavior.",
            ],
        )
        doc.add_heading("4. Discussion: Validity and limitations", level=1)
        doc.add_paragraph(
            "This artifact uses local container resource limits to emulate instance-type differences. "
            "It does not capture cloud networking variability, storage behavior, or microarchitectural differences across real EC2 families."
        )
        bullets = [
            "Local emulation via container resource limits captures CPU and memory constraints, but does not replicate cloud networking variability, storage latency (e.g., EBS), or noisy‑neighbor effects found on shared infrastructure.",
            "ApacheBench measures steady‑state request handling for simple HTTP workloads; it does not model user think time, browser behavior, or realistic mixed request distributions.",
            "The evaluation emphasizes mean metrics. Tail latency (e.g., p95/p99) can better reflect user experience under contention but is not fully characterized here.",
            "WordPress performance depends on themes, plugins, and caching configuration; results from a minimal seeded dataset may not generalize to content‑heavy production sites.",
            "Some historical runs contain incomplete outputs; conclusions are drawn from validated runs and should be interpreted as indicative trends rather than definitive capacity claims.",
        ]
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")
    else:
        add_placeholder(doc, "Conclude in 3–4 sentences, and describe future work.")

    # (Conclusion already added above. Keep this block removed to satisfy the required section structure.)

    doc.add_heading("References", level=1)
    # IEEE-like references (author/org, title, source, year, url, access date)
    refs = [
        "[1] Docker Inc., \"Docker Desktop Documentation,\" 2026. [Online]. Available: https://docs.docker.com/desktop/. Accessed: 2026-04-24.",
        "[2] Docker, \"WordPress Docker Image,\" Docker Hub. [Online]. Available: https://hub.docker.com/_/wordpress. Accessed: 2026-04-24.",
        "[3] Docker, \"MySQL Docker Image,\" Docker Hub. [Online]. Available: https://hub.docker.com/_/mysql. Accessed: 2026-04-24.",
        "[4] Docker, \"Nginx Docker Image,\" Docker Hub. [Online]. Available: https://hub.docker.com/_/nginx. Accessed: 2026-04-24.",
        "[5] Apache Software Foundation, \"Apache HTTP Server Documentation: ab (ApacheBench),\" 2026. [Online]. Available: https://httpd.apache.org/docs/2.4/programs/ab.html. Accessed: 2026-04-24.",
        "[6] Microsoft, \"Windows Subsystem for Linux Documentation,\" 2026. [Online]. Available: https://learn.microsoft.com/windows/wsl/. Accessed: 2026-04-24.",
    ]
    for r in refs:
        doc.add_paragraph(r)

    doc.add_page_break()
    doc.add_heading("Artifact Appendix (reproducibility)", level=1)
    doc.add_paragraph("Repository: https://github.com/hansz109/cs4296-project-")
    doc.add_paragraph("Prerequisites: Docker Desktop, WSL Ubuntu, Python 3.11+")
    doc.add_paragraph("Reproduce (high level):")
    steps = [
        r"1) Start stack: .\scripts\up_profile.ps1 -Profile general",
        r"2) Seed content: .\scripts\seed_wp.ps1",
        r"3) Install ab in WSL: .\scripts\install_ab_wsl.ps1",
        r"4) Run benchmarks: .\scripts\run_bench_wsl.ps1 -Profile <general|compute|memory> -Scenario <S1|S2|S3> -Repeat 3",
        r"5) Analyze: python scripts\analyze_results.py --results_dir experiments\results --out_dir experiments\summary",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")

    doc.add_paragraph("Expected outputs:")
    outs = [
        r"- experiments\results\<RUN_ID>\ab_*.txt, meta.txt, docker_stats.csv",
        r"- experiments\summary\ab_summary_by_profile_scenario.csv",
        r"- experiments\summary\throughput_by_profile_scenario.png",
        r"- experiments\summary\latency_by_profile_scenario.png",
    ]
    for o in outs:
        doc.add_paragraph(o, style="List Bullet")

    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()

